#!/usr/bin/env python3
"""
Create a comprehensive PRD for In-App Customer Messaging Inbox and export to Word.
"""

import sys
from pathlib import Path
from datetime import datetime

# Try to import python-docx, install if needed
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_prd_document():
    """Create a comprehensive PRD document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Product Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('In-App Customer Messaging Inbox', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Phase 1 - Q1 2026')
    doc.add_paragraph('')
    
    # Introduction
    doc.add_heading('📌 Introduction', 1)
    intro = doc.add_paragraph()
    intro.add_run('This PRD defines the requirements for building an ').bold = False
    intro.add_run('In-App Customer Messaging Inbox').bold = True
    intro.add_run(' to replace email as the primary asynchronous communication channel between customers and Tamara. This initiative will embed live chat conversations and support ticket updates in one unified, persistent experience within the Tamara mobile and web applications.')
    
    doc.add_paragraph('')
    doc.add_paragraph('The end-state user experience will be similar to Airbnb\'s Messages tab, which serves as a single inbox for all conversations (including support) with clear threads, search functionality, and a consistent place for customers to follow up on ongoing issues.')
    
    doc.add_paragraph('')
    doc.add_paragraph('This work aligns with and complements the Unified In-App Messaging and AI Sidekick initiative being developed by another team. This PRD focuses specifically on the customer-facing messaging inbox experience.')
    
    # Problem Statement
    doc.add_heading('❗ Problem Statement', 1)
    
    doc.add_paragraph('Currently, customer support communication at Tamara relies heavily on email, which creates several significant problems:')
    
    problems = [
        'Fragmented communication: Customers lose context when switching between email, in-app chat, and other channels',
        'Slow response times: Email-based support creates delays and requires customers to check multiple inboxes',
        'Poor thread tracking: Email threads are difficult to track across devices and sessions',
        'Lack of visibility: Customers cannot easily see the status of their support requests without sending follow-up emails',
        'Context loss: Each new email or chat session starts from scratch, requiring customers to re-explain their issues',
        'Mixed topics: Without clear thread separation, customers mix unrelated issues in single conversations',
        'No unified history: Support interactions, chatbot conversations, and ticket updates exist in separate systems'
    ]
    
    for problem in problems:
        p = doc.add_paragraph(problem, style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('These issues lead to:')
    impacts = [
        'Increased repeat contacts as customers lose track of previous conversations',
        'Higher customer effort to follow up on issues',
        'Reduced customer satisfaction due to perceived lack of transparency',
        'Lower chatbot containment rates as customers abandon chat and switch to email',
        'Inefficient support operations as agents spend time re-establishing context'
    ]
    
    for impact in impacts:
        p = doc.add_paragraph(impact, style='List Bullet')
    
    # Goals
    doc.add_heading('🎯 Goals', 1)
    
    doc.add_paragraph('The primary goals of this initiative are:')
    
    goals = [
        {
            'title': 'Replace Email as Primary Async Channel',
            'description': 'Make in-app messaging the default and preferred method for asynchronous customer support communication, reducing reliance on email by at least 60% within 6 months of launch.'
        },
        {
            'title': 'Unify Customer Communication Experience',
            'description': 'Provide a single, persistent inbox where customers can access all their conversations with Tamara, including chatbot interactions, agent conversations, and support ticket updates.'
        },
        {
            'title': 'Improve Customer Visibility and Transparency',
            'description': 'Give customers real-time visibility into the status of their support requests without requiring them to contact support again, reducing perceived wait times and increasing trust.'
        },
        {
            'title': 'Reduce Repeat Contacts',
            'description': 'Enable customers to easily find and continue previous conversations, reducing the need to re-explain issues and decreasing repeat contact rate by 30%.'
        },
        {
            'title': 'Increase Chatbot Containment',
            'description': 'Keep the entire support journey within a single thread, allowing chatbot to handle initial intents and seamlessly escalate to agents when needed, improving containment rates.'
        },
        {
            'title': 'Enable Multi-Threaded Conversations',
            'description': 'Allow customers to maintain separate conversation threads for different topics (e.g., refunds, order issues, account problems), preventing context mixing and improving organization.'
        }
    ]
    
    for goal in goals:
        p = doc.add_paragraph()
        p.add_run(goal['title']).bold = True
        doc.add_paragraph(goal['description'], style='List Bullet 2')
    
    # Success Metrics
    doc.add_heading('📊 Success Metrics (must align with CPX 2025)', 1)
    
    doc.add_paragraph('The following metrics will be tracked to measure success:')
    
    metrics = [
        {
            'metric': 'Email Support Volume Reduction',
            'target': '60% reduction in email-based support contacts within 6 months',
            'baseline': 'Current monthly email support volume',
            'measurement': 'Monthly email support ticket count'
        },
        {
            'metric': 'Repeat Contact Rate',
            'target': '30% reduction in repeat contacts for the same issue',
            'baseline': 'Current repeat contact rate',
            'measurement': 'Percentage of contacts that reference a previous interaction'
        },
        {
            'metric': 'Customer Satisfaction (CSAT)',
            'target': 'Increase CSAT for support interactions by 15%',
            'baseline': 'Current support CSAT score',
            'measurement': 'Post-interaction CSAT surveys'
        },
        {
            'metric': 'Chatbot Containment Rate',
            'target': 'Increase chatbot containment by 25%',
            'baseline': 'Current chatbot containment rate',
            'measurement': 'Percentage of chatbot conversations resolved without agent escalation'
        },
        {
            'metric': 'Time to Resolution Perception',
            'target': 'Improve customer perception of response time by 20%',
            'baseline': 'Current customer-reported wait time perception',
            'measurement': 'Customer surveys and feedback'
        },
        {
            'metric': 'In-App Messaging Adoption',
            'target': '80% of eligible customers use in-app messaging for at least one support interaction within 3 months',
            'baseline': '0% (new feature)',
            'measurement': 'Percentage of active customers who initiate at least one in-app message'
        },
        {
            'metric': 'Thread Continuity',
            'target': '70% of conversations continue in the same thread across multiple sessions',
            'baseline': 'N/A (new capability)',
            'measurement': 'Percentage of threads with multiple messages across different sessions'
        }
    ]
    
    for metric in metrics:
        p = doc.add_paragraph()
        p.add_run(metric['metric']).bold = True
        doc.add_paragraph(f"Target: {metric['target']}", style='List Bullet 2')
        doc.add_paragraph(f"Baseline: {metric['baseline']}", style='List Bullet 2')
        doc.add_paragraph(f"Measurement: {metric['measurement']}", style='List Bullet 2')
        doc.add_paragraph('')
    
    # User Personas
    doc.add_heading('👤 User Personas', 1)
    
    personas = [
        {
            'name': 'Frequent Support User',
            'description': 'Customers who regularly contact support for order issues, refunds, or account problems. They value quick responses and the ability to track multiple ongoing issues.',
            'needs': [
                'Quick access to all ongoing conversations',
                'Clear status updates on open issues',
                'Ability to attach documents and screenshots',
                'Search functionality to find previous conversations'
            ]
        },
        {
            'name': 'Occasional Support Seeker',
            'description': 'Customers who contact support infrequently but need help when they do. They may not be familiar with support processes and need clear guidance.',
            'needs': [
                'Easy-to-find support entry point',
                'Clear instructions on how to get help',
                'Reassurance that their issue is being tracked',
                'Simple interface that doesn\'t require learning'
            ]
        },
        {
            'name': 'Tech-Savvy Power User',
            'description': 'Customers comfortable with technology who expect modern, efficient communication channels. They prefer in-app experiences over email.',
            'needs': [
                'Fast, responsive interface',
                'Rich features like file sharing and formatting',
                'Keyboard shortcuts and efficiency features',
                'Integration with other app features'
            ]
        },
        {
            'name': 'Mobile-First User',
            'description': 'Customers who primarily or exclusively use mobile devices. They need a mobile-optimized experience that works well on small screens.',
            'needs': [
                'Mobile-optimized interface',
                'Push notifications for new messages',
                'Easy photo/document attachment from mobile',
                'Offline message queuing'
            ]
        }
    ]
    
    for persona in personas:
        p = doc.add_paragraph()
        p.add_run(persona['name']).bold = True
        doc.add_paragraph(persona['description'])
        doc.add_paragraph('Key Needs:')
        for need in persona['needs']:
            doc.add_paragraph(need, style='List Bullet 2')
        doc.add_paragraph('')
    
    # Solution
    doc.add_heading('💡 Solution', 1)
    
    doc.add_paragraph('Implement an In-App Customer Messaging Inbox (Phase 1) as a replacement for email-based customer support communication, centered around a persistent support inbox inside the Tamara app.')
    
    doc.add_heading('Core Components (Phase 1)', 2)
    
    components = [
        {
            'name': 'Support Inbox (Async Messaging)',
            'description': 'A dedicated inbox where customers can message Tamara directly from the app. Conversations are asynchronous by default, allowing replies over hours or days without losing context (unlike email). Each conversation persists as a long-lived thread rather than a one-off chat session.',
            'features': [
                'Dedicated "Messages" or "Support" tab in the app navigation',
                'Inbox view showing all conversation threads',
                'Thread list with preview of last message, timestamp, and status',
                'Unread message indicators and badges',
                'Persistent conversation history across app sessions',
                'Cross-device synchronization',
                'Push notifications for new messages'
            ]
        },
        {
            'name': 'Multi-Threaded Chat Conversations',
            'description': 'Every new chatbot conversation creates a new thread in the inbox (even if it never gets handed to an agent). Threads are titled automatically using AI, based on detected intent and key entities (e.g., "Refund status – Order 12345", "Card verification issue", "Payment not reflected"). This keeps conversations clean and searchable, and prevents customers from mixing unrelated topics in one long chat history.',
            'features': [
                'Automatic thread creation for each new conversation',
                'AI-powered thread titling based on intent and entities',
                'Manual thread creation option for customers',
                'Thread organization and categorization',
                'Thread search and filtering',
                'Thread archiving and deletion',
                'Thread status indicators (active, waiting, resolved)'
            ]
        },
        {
            'name': 'Ticket Updates Embedded in Threads',
            'description': 'When a support ticket is created, updated, or resolved, the customer sees system messages inside the relevant thread (e.g., "Ticket created", "Awaiting information", "Resolved"). This removes the need for status update emails and gives customers real-time visibility into progress without contacting support again.',
            'features': [
                'Automatic ticket creation from conversations',
                'Real-time ticket status updates in thread',
                'System messages for key ticket events',
                'Visual indicators for ticket status',
                'Link to full ticket details when needed',
                'Timeline view of ticket lifecycle',
                'Integration with existing ticketing system'
            ]
        }
    ]
    
    for component in components:
        p = doc.add_paragraph()
        p.add_run(component['name']).bold = True
        doc.add_paragraph(component['description'])
        doc.add_paragraph('Key Features:')
        for feature in component['features']:
            doc.add_paragraph(feature, style='List Bullet 2')
        doc.add_paragraph('')
    
    doc.add_paragraph('Phase 1 intentionally excludes advanced automation and focuses on clarity, continuity, and replacing email as the primary async channel.')
    
    # Solution Design
    doc.add_heading('📐 Solution Design', 1)
    
    doc.add_heading('User Experience Flow', 2)
    
    doc.add_paragraph('1. Customer opens the Tamara app and navigates to the "Messages" or "Support" tab')
    doc.add_paragraph('2. Inbox view displays all conversation threads, sorted by most recent activity')
    doc.add_paragraph('3. Customer can:')
    doc.add_paragraph('   - Tap an existing thread to continue the conversation', style='List Bullet 2')
    doc.add_paragraph('   - Tap "New Message" to start a new conversation', style='List Bullet 2')
    doc.add_paragraph('   - Use search to find previous conversations', style='List Bullet 2')
    doc.add_paragraph('4. When starting a new conversation, chatbot handles initial interaction')
    doc.add_paragraph('5. If escalation is needed, conversation seamlessly transitions to an agent')
    doc.add_paragraph('6. Ticket updates appear as system messages in the thread')
    doc.add_paragraph('7. Customer receives push notifications for new messages')
    doc.add_paragraph('8. Conversation persists across app sessions and devices')
    
    doc.add_heading('Technical Architecture', 2)
    
    doc.add_paragraph('The solution will integrate with:')
    doc.add_paragraph('- Existing chatbot infrastructure', style='List Bullet')
    doc.add_paragraph('- Current ticketing system (for status updates)', style='List Bullet')
    doc.add_paragraph('- Unified messaging platform (being built by another team)', style='List Bullet')
    doc.add_paragraph('- Push notification service', style='List Bullet')
    doc.add_paragraph('- User authentication and session management', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('Key technical requirements:')
    doc.add_paragraph('- Real-time message delivery and synchronization', style='List Bullet')
    doc.add_paragraph('- Message persistence and history storage', style='List Bullet')
    doc.add_paragraph('- Thread management and organization', style='List Bullet')
    doc.add_paragraph('- AI integration for thread titling', style='List Bullet')
    doc.add_paragraph('- Cross-platform compatibility (iOS, Android, Web)', style='List Bullet')
    doc.add_paragraph('- Offline message queuing', style='List Bullet')
    
    doc.add_heading('Integration Points', 2)
    
    doc.add_paragraph('This initiative must align with:')
    doc.add_paragraph('- Unified In-App Messaging and AI Sidekick PRD (reference: https://www.notion.so/tamaracom/PRD-Unified-in-app-messaging-AI-sidekick-for-Q1-2026-2aa62429127880769e6dc4e481985a5a)', style='List Bullet')
    doc.add_paragraph('- Existing chatbot and AI infrastructure', style='List Bullet')
    doc.add_paragraph('- Current support ticketing system', style='List Bullet')
    doc.add_paragraph('- Notification center (being built by another team)', style='List Bullet')
    
    # Inclusions and Exclusions
    doc.add_heading('✅ Inclusions and Exclusions', 1)
    
    doc.add_heading('Phase 1 Inclusions', 2)
    
    inclusions = [
        'Support inbox with thread list and conversation view',
        'Async messaging (send/receive messages over time)',
        'Multi-threaded conversations with AI-powered titling',
        'Ticket status updates embedded in threads',
        'Basic search functionality',
        'Push notifications for new messages',
        'Cross-device synchronization',
        'Integration with existing chatbot',
        'Integration with ticketing system for status updates',
        'Mobile and web support',
        'Message persistence and history'
    ]
    
    for inclusion in inclusions:
        doc.add_paragraph(inclusion, style='List Bullet')
    
    doc.add_heading('Phase 1 Exclusions', 2)
    
    exclusions = [
        'Advanced automation and workflows',
        'Rich media support beyond basic images and documents',
        'Voice or video messaging',
        'Group conversations or team collaboration',
        'Advanced analytics dashboard for customers',
        'Custom notification preferences (beyond basic on/off)',
        'Integration with external messaging platforms',
        'Proactive messaging from Tamara (beyond ticket updates)',
        'Advanced AI features beyond thread titling',
        'Customer-facing admin or moderation tools'
    ]
    
    for exclusion in exclusions:
        doc.add_paragraph(exclusion, style='List Bullet')
    
    # Roll-out Plan
    doc.add_heading('🚀 Roll-out Plan', 1)
    
    doc.add_heading('Phase 1 Timeline (Q1 2026)', 2)
    
    doc.add_paragraph('Week 1-2: Design and Architecture')
    doc.add_paragraph('- Finalize UI/UX designs', style='List Bullet 2')
    doc.add_paragraph('- Define API contracts and integration points', style='List Bullet 2')
    doc.add_paragraph('- Technical architecture review', style='List Bullet 2')
    
    doc.add_paragraph('')
    doc.add_paragraph('Week 3-6: Core Development')
    doc.add_paragraph('- Build inbox UI and thread management', style='List Bullet 2')
    doc.add_paragraph('- Implement async messaging infrastructure', style='List Bullet 2')
    doc.add_paragraph('- Integrate with chatbot and ticketing system', style='List Bullet 2')
    doc.add_paragraph('- Develop AI thread titling feature', style='List Bullet 2')
    
    doc.add_paragraph('')
    doc.add_paragraph('Week 7-8: Testing and Refinement')
    doc.add_paragraph('- Internal testing and bug fixes', style='List Bullet 2')
    doc.add_paragraph('- Integration testing with related systems', style='List Bullet 2')
    doc.add_paragraph('- Performance optimization', style='List Bullet 2')
    
    doc.add_paragraph('')
    doc.add_paragraph('Week 9-10: Beta Launch')
    doc.add_paragraph('- Limited beta release to 5% of users', style='List Bullet 2')
    doc.add_paragraph('- Monitor metrics and gather feedback', style='List Bullet 2')
    doc.add_paragraph('- Iterate based on beta learnings', style='List Bullet 2')
    
    doc.add_paragraph('')
    doc.add_paragraph('Week 11-12: Full Launch')
    doc.add_paragraph('- Gradual rollout to 100% of users', style='List Bullet 2')
    doc.add_paragraph('- Marketing and user education', style='List Bullet 2')
    doc.add_paragraph('- Monitor success metrics', style='List Bullet 2')
    
    doc.add_heading('Success Criteria for Launch', 2)
    
    doc.add_paragraph('Before moving to full launch, we must achieve:')
    doc.add_paragraph('- 70% of beta users successfully send at least one message', style='List Bullet')
    doc.add_paragraph('- Average response time under 2 seconds for message delivery', style='List Bullet')
    doc.add_paragraph('- Zero critical bugs or data loss incidents', style='List Bullet')
    doc.add_paragraph('- Positive feedback from 80% of beta users', style='List Bullet')
    doc.add_paragraph('- Successful integration with all required systems', style='List Bullet')
    
    # Hypothesis
    doc.add_heading('Hypothesis', 1)
    
    doc.add_paragraph('If we replace email with an in-app support inbox that supports persistent, multi-threaded conversations and surfaces ticket updates directly in the conversation, customers will have better visibility into their issues, leading to:')
    
    hypothesis_outcomes = [
        'Fewer repeat contacts (30% reduction)',
        'Lower effort to follow up on issues',
        'Improved customer satisfaction (15% increase in CSAT)',
        'Higher chatbot containment rates (25% increase)',
        'Reduced email support volume (60% reduction)',
        'Better perceived transparency and trust'
    ]
    
    for outcome in hypothesis_outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('This will be measured through the success metrics defined above, with data collection starting from beta launch.')
    
    # Add page break before appendix
    doc.add_page_break()
    
    # Appendix
    doc.add_heading('Appendix', 1)
    
    doc.add_heading('Related Documents', 2)
    doc.add_paragraph('Unified In-App Messaging and AI Sidekick PRD:')
    doc.add_paragraph('https://www.notion.so/tamaracom/PRD-Unified-in-app-messaging-AI-sidekick-for-Q1-2026-2aa62429127880769e6dc4e481985a5a', style='Intense Quote')
    
    doc.add_heading('Reference UX', 2)
    doc.add_paragraph('Airbnb Messages Tab:')
    doc.add_paragraph('https://www.airbnb.ae/resources/hosting-homes/a/getting-the-most-out-of-the-messages-tab-678', style='Intense Quote')
    
    return doc


if __name__ == "__main__":
    print("Creating comprehensive PRD document...")
    doc = create_prd_document()
    
    output_file = Path("PRD_In-App_Customer_Messaging_Inbox.docx")
    doc.save(str(output_file))
    
    print(f"\n✅ PRD created successfully!")
    print(f"📄 Saved to: {output_file.absolute()}")
    print(f"\nThe document includes:")
    print("  - Introduction and Problem Statement")
    print("  - Goals and Success Metrics")
    print("  - User Personas")
    print("  - Solution Design")
    print("  - Inclusions and Exclusions")
    print("  - Roll-out Plan")
    print("  - Hypothesis")



