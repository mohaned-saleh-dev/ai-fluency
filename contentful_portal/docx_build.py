"""
Word export matching the human-readable export structure.
"""
import io
import re
from datetime import datetime, timezone
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x0F, 0x17, 0x2A)  # deep ink blue
MUTED = RGBColor(0x3D, 0x4A, 0x5C)


def _add_locale_block(doc, heading: str, text: str, heading_level: int) -> None:
    if not (text and text.strip()):
        p = doc.add_paragraph("(No content in this language.)", style="Body Text 2")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = MUTED
        return
    h = doc.add_heading(heading, heading_level)
    h.runs[0].font.color.rgb = ACCENT
    segs = text.split("\n\n---\n\n")
    if not segs:
        return
    first = segs[0].strip()
    if first.startswith("FIELDS:"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r0 = p.add_run("Fields in this version: ")
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = MUTED
        p.add_run(first.replace("FIELDS: ", "")).font.size = Pt(10)
    for seg in segs[1:]:
        s = seg.strip()
        if not s:
            continue
        mline = s.split("\n", 1)
        head = (mline[0] or "").strip()
        body = (mline[1] or "").lstrip() if len(mline) > 1 else ""
        m = re.match(r"^(\d+)\.\s+(.+)$", head)
        if m:
            n, name = m.group(1), m.group(2).strip()
            sub = doc.add_paragraph()
            r = sub.add_run(f"{n}. {name}")
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x16, 0x21, 0x2E)
            sub.paragraph_format.space_after = Pt(2)
            sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for bl in re.split(r"\n\s*\n", body) if body else []:
                t = (bl or "").strip()
                if t:
                    bp = doc.add_paragraph(t)
                    bp.paragraph_format.space_after = Pt(4)
                    bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for run in bp.runs:
                        run.font.size = Pt(11)
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(s)
            p.paragraph_format.space_after = Pt(4)


def build_docx(
    rows: List[dict],
    locales: List[str],
) -> bytes:
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = section.top_margin
        # keep defaults

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    tr = title.add_run("Contentful article downloader")
    tr.font.size = Pt(32)
    tr.font.bold = True
    tr.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    sr = sub.add_run("Human-readable export for the content team")
    sr.font.size = Pt(13)
    sr.font.color.rgb = MUTED
    sr.italic = True

    when = doc.add_paragraph()
    when.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w = when.add_run(
        datetime.now(timezone.utc).strftime("Generated: %Y-%m-%d %H:%M UTC")
    )
    w.font.size = Pt(9)
    w.font.color.rgb = MUTED
    when.paragraph_format.space_after = Pt(24)

    p_intro = doc.add_paragraph(
        f"This document contains {len(rows)} entries from your Contentful space. "
        "Each section lists metadata, then the same field layout as the companion CSV, "
        "in English and Arabic (where available)."
    )
    p_intro.paragraph_format.line_spacing = 1.12
    p_intro.paragraph_format.space_after = Pt(18)
    p_intro.runs[0].font.size = Pt(10)
    p_intro.runs[0].font.color.rgb = MUTED

    doc.add_page_break()

    for i, row in enumerate(rows, 1):
        if i > 1:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(20)
            sep.paragraph_format.space_after = Pt(4)
            sr = sep.add_run("─" * 52)
            sr.font.color.rgb = MUTED
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(16)
        meta_title = (row.get("title_lookup") or "Untitled").strip() or "(no title)"
        tname = (row.get("content_type_name") or row.get("content_type") or "").strip()

        h1 = doc.add_heading(f"{i}. {meta_title}", level=1)
        h1.runs[0].font.size = Pt(20)
        h1.runs[0].font.color.rgb = ACCENT
        h1.runs[0].font.name = "Calibri Light"
        h1.paragraph_format.space_after = Pt(3)

        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(6)
        bits = [
            f"ID: {row.get('entry_id', '')}",
            f"Type: {tname}",
        ]
        if row.get("updated_at"):
            bits.append(f"Updated: {row['updated_at']}")
        mr = meta.add_run("  •  ".join(bits))
        mr.font.size = Pt(9)
        mr.font.color.rgb = MUTED

        for loc in locales:
            loc_label = "English (US)" if loc == "en-US" else ("العربية" if loc == "ar" else loc)
            content = (row.get(f"all_fields_{loc}") or "").strip()
            _add_locale_block(doc, loc_label, content, 2)

        links = (row.get("all_links_deduped") or "").strip()
        if links:
            sh = doc.add_heading("All links in this entry", 3)
            sh.runs[0].font.color.rgb = MUTED
            for line in links.split("\n"):
                line = line.strip()
                if not line:
                    continue
                lp = doc.add_paragraph(line, style="List Bullet")
                if lp.runs:
                    lp.runs[0].font.size = Pt(9)
                    lp.runs[0].font.color.rgb = RGBColor(0x0B, 0x57, 0x9E)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
