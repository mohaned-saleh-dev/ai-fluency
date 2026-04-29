from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

TEAL = RGBColor(0x00, 0x89, 0x7B)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "F0F7F6"
TEAL_HEX = "00897B"
DARK_HEX = "1A1A2E"


def add_title_page():
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("The AI Chatbot Story")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("All-Hands Presentation")
    run.font.size = Pt(18)
    run.font.color.rgb = TEAL

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Speaker: Mohaned  ·  Audience: All of Tamara  ·  Duration: 5:00")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    doc.add_page_break()


def add_section_header(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK if level == 1 else TEAL
    return h


def add_timing_banner(text):
    p = doc.add_paragraph()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TEAL_HEX}"/>')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = WHITE
    p.runs[0].element.rPr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{TEAL_HEX}"/>')
    )


def add_script_block(text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line.strip())
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        run.italic = True


def add_stage_direction(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"[ {text} ]")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, DARK_HEX)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()


def add_placeholder_table():
    headers = ["Placeholder", "Description", "Where Used"]
    rows = [
        ["[X] chats/day", "Current daily chat volume", "Script §2, Slide 4, Video"],
        ["[X% YoY]", "Contact volume growth rate", "Slide 4"],
        ["[X%] deflection", "Current chatbot deflection rate", "Script §2, Slide 5"],
        ["[X] CSAT", "Current CSAT score for chatbot", "Script §2, Slide 5"],
        ["[X sec] response", "Current first response time", "Script §2, Slides 3 & 5"],
        ["[X min] old response", "Previous response time (pre-AI)", "Script §2, Slide 3"],
        ["[X%] resolution", "% resolved without human", "Script §2, Slide 5"],
        ["[X%] contact rate (no AI)", "Projected contact rate without chatbot", "Script §3, Slide 6"],
        ["[X%] contact rate (AI)", "Current contact rate with chatbot", "Script §3, Slide 6"],
        ["[X] advisors (no AI)", "Headcount needed without chatbot", "Script §3, Slide 6"],
        ["[X] advisors (current)", "Current headcount with chatbot", "Script §3, Slide 6"],
        ["[X FTEs] delta", "Headcount savings", "Script §3, Slide 6"],
        ["[₿X / $X] cost", "Cost comparison & annual savings", "Script §3, Slide 6"],
        ["[Date]", "Presentation date", "Slide 1"],
    ]
    add_table(headers, rows, col_widths=[1.8, 2.5, 2.0])


# ── BUILD DOCUMENT ──

add_title_page()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PART A — SPEAKER SCRIPT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_section_header("PART A — SPEAKER SCRIPT", level=1)
add_body("Timed to exactly 5 minutes. Italicized text is the spoken script. Stage directions are in brackets.")

doc.add_paragraph()

# Section 1
add_timing_banner("⏱  0:00 – 1:00  |  SECTION 1: THE HORIZON (The Hook)")
doc.add_paragraph()

add_script_block("""
A few years ago, a wave hit the financial world. Fintechs like Tamara disrupted banks — not by fighting them, but by reimagining what was possible. We made ownership accessible. We gave people the power to own their dreams.

Today, a new wave is here. AI. And just like before, people are asking the same question: "Should I be worried?"

Here's my answer: disruption is never the threat — standing still is.

Our North Star hasn't changed. We still want people to own their dreams. But now, we have a new ability: we can let our people — you — do the dreaming, the creative thinking, and the problem-solving. And AI? AI handles the repetitive, the mundane, the things that slow us down.

So no — AI isn't replacing us. It's filtering out the boring stuff so we can focus on what actually matters.
""")

add_stage_direction("PAUSE — let that land")

doc.add_paragraph()

# Section 2
add_timing_banner("⏱  1:00 – 4:00  |  SECTION 2: THE TRANSFORMATION (The Proof)")
doc.add_paragraph()

add_section_header("The Before & After Story  (1:00 – 1:45)", level=2)

add_script_block("""
Let me paint a picture.

Before: A partner signs up with Tamara. They have questions — "Where do I upload my documents?" "Why was my application rejected?" "How do I activate?" Every single one of those questions went straight to a human agent. The partner waited an average of fifteen minutes just for a first response. And on the agent's side? They were buried — answering the same five questions, hundreds of times a day.

After: That same partner opens the chat. Within five seconds, our AI chatbot understands the question, pulls the right answer, and resolves it — no human needed. The partner moves on. The agent? They're now spending their time on real problems — the edge cases, the complex disputes, the moments where a human touch actually makes a difference.

That's not a small shift. That's a transformation.
""")

add_section_header("The Problem & The Numbers  (1:45 – 2:30)", level=2)

add_script_block("""
Here's the reality we were facing. As Tamara grows, contact volume grows with it. More merchants, more consumers, more chats. We were processing [X] chats per day, and that number was only going up. Without a change, we'd be in a constant hiring race just to keep up.

So we asked: what if AI could carry the weight of the routine, so our people could carry the weight of the meaningful?
""")

add_body("Here's where we are today:")

add_table(
    ["Metric", "Value"],
    [
        ["Deflection Rate", "[X%]"],
        ["CSAT", "[X]"],
        ["First Response Time", "< [X] seconds"],
        ["Issue Resolution (no human)", "[X%]"],
    ],
    col_widths=[3.0, 3.0],
)

add_script_block("These aren't vanity metrics. This is proof that efficiency isn't the goal — it's the enabler.")

add_section_header('The "What If Without" — The Counterfactual  (2:30 – 3:15)', level=2)

add_script_block("Now let me show you the other side. What if we hadn't built this?")

add_body("With our current growth trajectory:")

add_table(
    ["Metric", "Without AI", "With AI"],
    [
        ["Contact Rate", "[X%]", "[X%]"],
        ["Advisors Needed", "[X]", "[X]"],
        ["Additional Headcount", "[X FTEs]", "—"],
        ["Cost to Serve", "[₿X]", "[₿X]"],
        ["Annual Savings", "—", "[₿X / $X]"],
    ],
    col_widths=[2.2, 2.0, 2.0],
)

add_script_block(
    "The chatbot isn't just a product. It's a business decision that's saving us [X] "
    "and letting us scale without scaling costs linearly."
)

add_section_header("Three AI Chat Agents  (3:15 – 3:45)", level=2)

add_script_block("Today, we don't have one chatbot. We have three purpose-built AI chat agents:")

add_table(
    ["Customer AI Chatbot", "Partner AI Chatbot", "Partner Onboarding AI Chatbot"],
    [
        [
            "Consumer journey: orders, payments, returns, accounts",
            "Live merchants: settlements, activations, integrations",
            "Pre-onboarding: documents, applications, activation",
        ]
    ],
    col_widths=[2.2, 2.2, 2.2],
)

add_script_block("""
Each one is trained on its own knowledge base, tuned for its own intents, and designed for its own user. One size does not fit all — and we didn't try to force it.

And yes — we've built this with compliance in mind. PII masking, data residency, and privacy protections are baked in, not bolted on.
""")

add_section_header("What's Next — The Roadmap Teaser  (3:45 – 4:00)", level=2)

add_script_block("But we're not stopping at chat. Here's a glimpse of where we're headed:")

add_table(
    ["Now → Next", "Future — Agentic AI"],
    [
        ["Chat SDKs — native in-app experience", "Voice / Phone AI"],
        ["AI-Powered Search & Summary", "Dispute Handling AI Agent"],
        ["Quality & Intelligence Layer", "Care Advisor Co-pilot (Salesforce)"],
    ],
    col_widths=[3.2, 3.2],
)

add_script_block("This is the beginning, not the end.")

doc.add_paragraph()

# Section 3
add_timing_banner("⏱  4:00 – 5:00  |  SECTION 3: THE REVEAL (The Video)")
doc.add_paragraph()

add_script_block("""
I could keep talking — but I'd rather show you.

What you're about to see is our AI chatbot in action. A real conversation. A real resolution. No scripts, no staging. Just the product, working.

Let's watch.
""")

add_stage_direction("▶  PLAY VIDEO — ~45–50 seconds")

doc.add_paragraph()

add_script_block("""
That's not a demo of the future. That's today. And it's only going to get better.

Thank you.
""")

doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PART B — SLIDE OUTLINE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_section_header("PART B — SLIDE-BY-SLIDE OUTLINE", level=1)
add_body("10 slides mapped to the script timing. Visual direction included for the deck designer.")

doc.add_paragraph()

slides = [
    {
        "num": 1,
        "title": "Title Slide",
        "timing": "—",
        "visual": "Bold title, Tamara branding, AI-themed visual (abstract wave motif)",
        "content": [
            "Title: The AI Chatbot Story",
            "Subtitle: All-Hands · [Date]",
            "Speaker: Mohaned",
        ],
    },
    {
        "num": 2,
        "title": "The Hook — Disruption Isn't New",
        "timing": "0:00 – 1:00",
        "visual": "Split image — left: traditional bank, right: Tamara app. Wave graphic overlay.",
        "content": [
            '"Fintechs disrupted banks. Now AI is disrupting work."',
            '"Disruption isn\'t the threat — standing still is."',
            '"Our North Star hasn\'t changed. AI just filters out the boring stuff."',
        ],
    },
    {
        "num": 3,
        "title": "Before & After",
        "timing": "1:00 – 1:45",
        "visual": "Two-column layout. LEFT = Before (red/muted), RIGHT = After (green/vibrant).",
        "content": [
            "Before: 100% queries → human, ~15 min wait, agents buried in repetition",
            "After: AI-first, < [X]s response, agents on complex cases, scale ≠ headcount",
        ],
    },
    {
        "num": 4,
        "title": "The Problem — Growing Volume",
        "timing": "1:45 – 2:00",
        "visual": "Line chart: contact volume growth vs. business growth (upward trend).",
        "content": [
            "Daily chats: [X]",
            "Projected growth: [X% YoY]",
            '"Without AI, we\'d be in a constant hiring race."',
        ],
    },
    {
        "num": 5,
        "title": "The Numbers — Where We Are Today",
        "timing": "2:00 – 2:30",
        "visual": "4 large KPI cards, clean and bold.",
        "content": [
            "Deflection Rate: [X%]",
            "CSAT: [X]",
            "First Response Time: < [X] sec",
            "Issue Resolution (no human): [X%]",
            '"Efficiency isn\'t the goal — it\'s the enabler."',
        ],
    },
    {
        "num": 6,
        "title": "What If Without AI?",
        "timing": "2:30 – 3:15",
        "visual": 'Side-by-side "With AI" vs. "Without AI" in contrasting colors.',
        "content": [
            "Contact Rate, Advisors Needed, Additional Headcount, Cost to Serve, Annual Savings",
            "All with [placeholder] values — fill in from counterfactual data",
        ],
    },
    {
        "num": 7,
        "title": "Three AI Chat Agents",
        "timing": "3:15 – 3:45",
        "visual": "Three cards/columns with icons.",
        "content": [
            "Customer AI Chatbot — consumer journey",
            "Partner AI Chatbot — live merchant ops",
            "Partner Onboarding AI Chatbot — pre-onboarding",
            "Footer: Built with compliance — PII masking · data residency · privacy by design",
        ],
    },
    {
        "num": 8,
        "title": "What's Next — The Roadmap",
        "timing": "3:45 – 4:00",
        "visual": 'Horizontal timeline: "Now → Next → Future".',
        "content": [
            "Now → Next: Chat SDKs, AI Search & Summary, Quality Layer",
            "Future (Agentic AI): Voice AI, Dispute AI Agent, Care Advisor Co-pilot",
        ],
    },
    {
        "num": 9,
        "title": "The Reveal — See It In Action",
        "timing": "4:00 – 4:50",
        "visual": "Full-screen video embed or large play button with chatbot UI screenshot.",
        "content": ['"No scripts. No staging. Just the product, working."', "▶ PLAY VIDEO (~45–50 sec)"],
    },
    {
        "num": 10,
        "title": "Closing",
        "timing": "4:50 – 5:00",
        "visual": "Tamara branding, clean and simple.",
        "content": [
            '"That\'s not a demo of the future. That\'s today."',
            '"Thank you."',
        ],
    },
]

for s in slides:
    h = add_section_header(f"Slide {s['num']}: {s['title']}", level=2)

    timing_p = doc.add_paragraph()
    run = timing_p.add_run(f"Timing: {s['timing']}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL

    visual_p = doc.add_paragraph()
    run = visual_p.add_run("Visual: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = visual_p.add_run(s["visual"])
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    for item in s["content"]:
        bp = doc.add_paragraph(item, style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PART C — VIDEO SCRIPT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_section_header("PART C — VIDEO SCRIPT / STORYBOARD", level=1)

add_body("Duration: 45–50 seconds  |  Format: Screen recording with voiceover or text captions  |  Tone: Clean, confident, no fluff")

doc.add_paragraph()

scenes = [
    {
        "num": 1,
        "time": "0:00 – 0:05",
        "title": "Opening",
        "visual": "Phone screen / web browser opens Tamara app or partner portal.",
        "vo": '"Meet Tamara\'s AI Chatbot."',
    },
    {
        "num": 2,
        "time": "0:05 – 0:15",
        "title": "The Question",
        "visual": 'User types a real question (e.g., "Where is my refund?" or "How do I upload my trade license?")',
        "vo": '"A customer asks a question — just like they always do."',
    },
    {
        "num": 3,
        "time": "0:15 – 0:25",
        "title": "The Response",
        "visual": "Bot responds instantly with the correct, contextual answer. Show the <5 second response time.",
        "vo": '"In under five seconds, they have their answer. No wait. No transfer. No frustration."',
    },
    {
        "num": 4,
        "time": "0:25 – 0:35",
        "title": "The Resolution",
        "visual": "User confirms resolution / gives thumbs up / conversation closes naturally.",
        "vo": '"Resolved. First try. And the agent? Free to focus on what matters most."',
    },
    {
        "num": 5,
        "time": "0:35 – 0:45",
        "title": "The Scale",
        "visual": "Pull back to show volume — a dashboard or counter showing thousands of conversations.",
        "vo": '"This happens [X] thousand times a day. Every day."',
    },
    {
        "num": 6,
        "time": "0:45 – 0:50",
        "title": "Close",
        "visual": "Tamara logo + tagline.",
        "vo": '"Tamara AI. Own your dreams — we\'ll handle the rest."',
    },
]

for sc in scenes:
    h = add_section_header(f"Scene {sc['num']}: {sc['title']}  ({sc['time']})", level=2)

    vp = doc.add_paragraph()
    run = vp.add_run("Visual: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = vp.add_run(sc["visual"])
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    vop = doc.add_paragraph()
    run = vop.add_run("Voiceover / Caption: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = vop.add_run(sc["vo"])
    run2.font.size = Pt(10)
    run2.italic = True

    doc.add_paragraph()

doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PLACEHOLDER REFERENCE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_section_header("PLACEHOLDER REFERENCE", level=1)
add_body("All placeholders to fill before the presentation:")
doc.add_paragraph()
add_placeholder_table()

# ── SAVE ──
output_path = "/Users/mohaned.saleh/prd-ticket-agent/All_Hands_AI_Chatbot_Story.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
