#!/usr/bin/env python3
"""
Generate Sarj Voice AI POC requirements document (Word).
"""
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        _set_cell_shading(hdr[i], "E8EEF5")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
            for p in table.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def main() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    accent = RGBColor(0x0F, 0x17, 0x2A)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("Voice AI Proof of Concept (POC)")
    t.bold = True
    t.font.size = Pt(22)
    t.font.color.rgb = accent
    title.paragraph_format.space_after = Pt(6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Technical & Functional Requirements — Partner: Sarj")
    s.font.size = Pt(14)
    s.font.color.rgb = RGBColor(0x3D, 0x4A, 0x5C)
    sub.paragraph_format.space_after = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Document date: {date.today().isoformat()}\n"
        "Classification: Confidential — Internal / Partner under NDA"
    ).font.size = Pt(9)
    meta.paragraph_format.space_after = Pt(18)

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This document defines the scope, data inputs, customer journey, and API "
        "integration requirements for a Voice AI proof of concept (POC) between "
        "Tamara and Sarj. The POC replaces a traditional IVR path with a conversational "
        "AI Voice Agent focused on post-purchase support, specifically refund-related "
        "questions once the customer and their recent orders are known."
    )

    doc.add_heading("2. POC objectives", level=1)
    doc.add_paragraph(
        "Demonstrate that Sarj’s Voice AI can:",
        style="List Bullet",
    )
    objectives = [
        "Authenticate the caller using the phone number associated with their Tamara account.",
        "Retrieve and reason over the customer’s recent order history (up to five orders).",
        "When refunds exist, answer natural-language questions about refund status and policy "
        "using (a) structured refund data from Tamara APIs and (b) Tamara-provided FAQ knowledge.",
        "Provide accurate, consistent answers on refund method, process, expected timing (ETA), "
        "and destination of funds where applicable.",
    ]
    for o in objectives:
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("3. Knowledge base (FAQ articles — KSA)", level=1)
    doc.add_paragraph(
        "Tamara will supply a CSV file containing the current customer FAQ articles "
        "relevant to the Kingdom of Saudi Arabia (KSA). Sarj shall ingest this content "
        "into the Voice AI knowledge layer used for retrieval and grounding."
    )
    doc.add_paragraph("Expected characteristics (to align during onboarding):", style="List Bullet")
    for line in [
        "Locale/market: KSA (content may be Arabic and/or English — exact columns to be confirmed with the file schema).",
        "Article identifiers and versioning so updates can be refreshed during the POC window.",
        "Clear separation between general policy text and order-specific facts (order-specific data must come from APIs, not KB alone).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4. Customer journey & conversation flow", level=1)
    doc.add_heading("4.1 High-level flow", level=2)
    steps = [
        "Greeting & intent: the Voice AI identifies that the user needs help with an order/refund.",
        "Phone capture: the flow prompts the customer to provide the phone number associated with their Tamara account "
        "(spoken and/or keypad entry — implementation detail for Sarj within telephony constraints).",
        "Identification: Sarj calls Tamara’s customer identification API using the captured phone number.",
        "Order retrieval: upon successful identification, Sarj calls Tamara’s orders API to fetch up to the five most recent orders.",
        "Refund path: for orders that have refund activity, Sarj calls Tamara’s refund-details API to obtain structured refund information.",
        "Q&A: the Voice AI answers the customer’s questions using API-grounded facts plus FAQ retrieval where appropriate.",
        "Escalation: if the customer cannot be identified, APIs fail, or the query is out of scope, the agent follows an agreed fallback (retry, alternate verification if allowed, or human handoff — to be confirmed).",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"Step {i}: ").bold = True
        p.add_run(step)

    doc.add_heading("4.2 Identifier: phone number", level=2)
    doc.add_paragraph(
        "The POC assumes the customer enters or confirms the phone number tied to their "
        "Tamara account. Format normalization (e.g. country code, leading zero) shall be "
        "handled consistently between Sarj’s capture layer and Tamara’s identification API contract."
    )

    doc.add_heading("4.3 Order list constraint", level=2)
    doc.add_paragraph(
        "After identification, the IVR/Voice flow shall fetch at most the five (5) most recent "
        "orders for that customer. Presentation order (e.g. newest first) should match API ordering "
        "unless otherwise specified in the API documentation Tamara provides."
    )

    doc.add_heading("5. Refund use case — required behaviors", level=1)
    doc.add_paragraph(
        "When any of the retrieved orders has refund information, the Voice AI must be able to "
        "answer questions about that refund, including:"
    )
    for item in [
        "Refund method (e.g. how the refund is processed).",
        "Process steps at a high level (what happens after approval).",
        "ETA / timing expectations (aligned with API status and FAQ where policy applies).",
        "Where the refund was issued (destination), when applicable and returned by the API.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Grounding rule: structured, order-specific refund facts come from Tamara APIs; "
        "the FAQ CSV supplements policy explanations and generic guidance. The agent must not "
        "contradict API facts with FAQ text. If FAQ and API conflict, API data wins for "
        "order-specific attributes."
    )

    doc.add_heading("6. Tamara API access for the POC (three endpoints)", level=1)
    doc.add_paragraph(
        "Tamara will expose three endpoints for Sarj to integrate against in the agreed POC "
        "environment (e.g. staging). Exact URLs, authentication scheme, and payload schemas "
        "will be provided in Tamara’s API pack; below are the functional requirements."
    )

    doc.add_heading("6.1 Endpoint A — Identify customer by phone number", level=2)
    doc.add_paragraph(
        "Purpose: resolve the Tamara customer record from the phone number provided during the call."
    )
    _add_table(
        doc,
        ["Aspect", "Requirement"],
        [
            ["Method / style", "REST (HTTPS). Method TBD in OpenAPI (typically POST or GET with query)."],
            ["Input", "Phone number (normalized per contract); optional correlation ID for support."],
            ["Output", "Customer identifier(s) required for subsequent calls; error codes for not found / ambiguous / invalid."],
            ["Security", "Mutual TLS or OAuth2 client credentials (per Tamara standard); IP allowlist if used."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("6.2 Endpoint B — Fetch customer orders (up to 5)", level=2)
    doc.add_paragraph(
        "Purpose: return up to the five most recent orders for the identified customer."
    )
    _add_table(
        doc,
        ["Aspect", "Requirement"],
        [
            ["Input", "Customer ID from Endpoint A (or token representing the session)."],
            ["Output", "List of orders (max 5): order ID, dates, status, amount cues as available — fields per schema."],
            ["Ordering", "Most recent first unless Tamara specifies otherwise in the schema."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("6.3 Endpoint C — Fetch order refunds (if any)", level=2)
    doc.add_paragraph(
        "Purpose: for a given order, return refund line items / events so the Voice AI can "
        "answer detailed refund questions."
    )
    _add_table(
        doc,
        ["Aspect", "Requirement"],
        [
            ["Input", "Customer context + Order ID (per Tamara auth model)."],
            ["Output", "Refund records: status, method, amounts, destination/wallet/card tail if applicable, timestamps, ETA fields if exposed."],
            ["Empty case", "Explicit empty response when no refunds exist for the order."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("7. Data handling, privacy, and compliance", level=1)
    for line in [
        "POC should use synthetic or staging data where possible; production access only if explicitly approved.",
        "Phone numbers, customer IDs, order and refund data are sensitive — logging must avoid storing full PANs or secrets; follow Tamara’s data classification rules.",
        "KSA market content in the FAQ must be used in line with Tamara’s brand and legal approvals.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("8. Non-functional requirements", level=1)
    _add_table(
        doc,
        ["Area", "POC expectation"],
        [
            ["Latency", "Voice turns should stay within Sarj’s recommended latency budget; APIs should respond within agreed SLOs (to be measured in staging)."],
            ["Languages", "Arabic and/or English support per Tamara’s POC decision; ASR/TTS quality validated on pilot phrases."],
            ["Availability", "Staging maintenance windows communicated in advance."],
            ["Observability", "Correlation IDs across identify → orders → refunds for joint debugging."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("9. Deliverables from Tamara", level=1)
    for d in [
        "CSV export of KSA FAQ articles for knowledge ingestion.",
        "API documentation (OpenAPI/Swagger or equivalent) for the three endpoints.",
        "POC credentials and environment base URLs; rate limits and support contacts.",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("10. Deliverables from Sarj", level=1)
    for d in [
        "Voice AI implementation of the agreed flow (phone capture → identify → orders → refunds → Q&A).",
        "Integration with the three Tamara endpoints per this document.",
        "Test plan and joint test windows; sample transcripts or call logs (redacted) for review.",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("11. Acceptance criteria (POC sign-off)", level=1)
    for d in [
        "Successful end-to-end calls for scripted scenarios: identified customer with orders; orders with and without refunds.",
        "Accurate answers on refund method, process, ETA, and destination when present in API (+ consistent FAQ tone).",
        "Documented list of gaps (latency, ASR errors, API edge cases) and recommendation for next phase.",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("12. Open items & assumptions", level=1)
    doc.add_paragraph(
        "The following are explicitly open for joint confirmation before kickoff:"
    )
    for d in [
        "Exact phone number format and normalization rules for KSA numbers.",
        "Authentication mechanism and token lifetime for the three APIs.",
        "Whether keypad (DTMF), voice capture, or both are required for phone entry.",
        "Escalation path to a human agent and any CRM ticketing integration (in or out of POC scope).",
        "CSV schema for FAQ columns (title, body, locale, URL, updated_at, etc.).",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("13. Appendix A — Call flow sequence (logical)", level=1)
    doc.add_paragraph(
        "Logical dependency chain for Sarj integration (each step depends on success of the prior):"
    )
    seq = [
        "Voice session starts → collect phone number from caller.",
        "POST/GET Endpoint A (identify customer) → receive customer_id (or error).",
        "GET Endpoint B (orders for customer_id) → receive 0–5 orders.",
        "For each order selected by the conversation (or each order with refund flags in list response), "
        "GET Endpoint C (refunds for order_id) → receive 0..n refund records.",
        "Compose spoken responses using API fields; enrich with FAQ retrieval for generic policy wording.",
    ]
    for i, line in enumerate(seq, 1):
        doc.add_paragraph(f"{i}. {line}")

    doc.add_heading("14. Appendix B — Indicative data elements (non-binding)", level=1)
    doc.add_paragraph(
        "Final field names and types will match Tamara’s published OpenAPI. The lists below are "
        "for planning discussions only."
    )
    doc.add_heading("14.1 Identify customer response (illustrative)", level=2)
    _add_table(
        doc,
        ["Element", "Notes"],
        [
            ["customer_id", "Stable identifier for subsequent API calls."],
            ["account_status", "If exposed — e.g. active / restricted (TBD)."],
            ["metadata", "Optional flags Sarj must ignore if not documented."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("14.2 Orders list item (illustrative)", level=2)
    _add_table(
        doc,
        ["Element", "Notes"],
        [
            ["order_id", "Unique order reference."],
            ["created_at / updated_at", "For ordering and customer explanations."],
            ["order_status", "High-level state visible to the customer."],
            ["has_refund / refund_summary", "If provided at list level to avoid unnecessary Endpoint C calls (TBD)."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("14.3 Refund record (illustrative)", level=2)
    _add_table(
        doc,
        ["Element", "Notes"],
        [
            ["refund_id / event_id", "Support correlation."],
            ["status", "e.g. pending, processing, completed, failed (exact enum per API)."],
            ["method", "Maps to “how” the refund is processed (customer-facing label from API + FAQ)."],
            ["amount / currency", "As returned by Tamara."],
            ["destination", "Card tail, wallet, bank reference — per Tamara policy on what is spoken aloud."],
            ["eta_or_expected_completion", "If available as a field or derivable from status + policy."],
            ["timestamps", "Requested, processed, completed — per schema."],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("15. Appendix C — Sample scenario matrix (for joint testing)", level=1)
    _add_table(
        doc,
        ["ID", "Scenario", "Expected"],
        [
            ["S1", "Valid phone, 3 orders, one order has refund", "Agent cites refund details from Endpoint C + policy from FAQ."],
            ["S2", "Valid phone, orders, no refunds", "Agent states no refund on file; FAQ only for generic refund policy."],
            ["S3", "Phone not found", "Graceful error; no order/refund APIs called."],
            ["S4", "Customer asks ETA only", "Answer from refund record + FAQ if status is generic."],
            ["S5", "FAQ vs API tension", "API-specific fields override FAQ text for that order."],
        ],
    )
    doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph(
        "End of document — questions should be routed through the Tamara POC owner."
    )
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    out_path = "/Users/mohaned.saleh/prd-ticket-agent/Sarj_Voice_AI_POC_Requirements.docx"
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
