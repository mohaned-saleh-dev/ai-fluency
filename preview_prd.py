#!/usr/bin/env python3
"""
Preview the PRD document content.
"""

from docx import Document
from pathlib import Path

def preview_document():
    """Read and display the PRD document."""
    doc_path = Path("PRD_In-App_Customer_Messaging_Inbox.docx")
    
    if not doc_path.exists():
        print(f"❌ Document not found: {doc_path}")
        return
    
    doc = Document(str(doc_path))
    
    print("=" * 80)
    print("PRD: IN-APP CUSTOMER MESSAGING INBOX")
    print("=" * 80)
    print()
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        # Check if it's a heading
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
            indent = "  " * (level - 1)
            print(f"\n{indent}{'#' * level} {text}")
        else:
            # Regular paragraph
            print(text)
            print()
    
    # Also check for tables
    for table in doc.tables:
        print("\n" + "-" * 80)
        print("TABLE:")
        print("-" * 80)
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            print(row_text)
        print("-" * 80)
        print()

if __name__ == "__main__":
    preview_document()



