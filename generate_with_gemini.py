#!/usr/bin/env python3
"""
Generate comprehensive PRD using Gemini, incorporating all requirements and the related PRD.
"""

import asyncio
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from prd_ticket_agent.integrations.gemini_client import GeminiClient
from prd_ticket_agent.config import load_context_from_env
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


async def generate_comprehensive_sections_with_gemini(gemini_client):
    """Use Gemini to generate comprehensive PRD sections."""
    
    # Related PRD URL content (we'll reference it)
    related_prd_context = """
    Related PRD: Unified In-App Messaging and AI Sidekick for Q1 2026
    URL: https://www.notion.so/tamaracom/PRD-Unified-in-app-messaging-AI-sidekick-for-Q1-2026-2aa62429127880769e6dc4e481985a5a
    
    This PRD focuses on building a unified messaging platform with AI sidekick capabilities.
    Our PRD (In-App Customer Messaging Inbox) is the customer-facing component that integrates with this platform.
    """
    
    # Generate comprehensive technical specifications
    tech_specs_prompt = f"""
    Based on the following context, generate comprehensive technical specifications for an In-App Customer Messaging Inbox system.

    Context:
    {related_prd_context}
    
    Requirements:
    - Replace email as primary async communication channel
    - Support multi-threaded conversations
    - Integrate with chatbot and ticketing system
    - Real-time message delivery
    - Cross-platform (iOS, Android, Web)
    - AI-powered thread titling
    - Ticket status updates embedded in threads
    
    Generate detailed technical specifications including:
    1. API endpoints with request/response schemas
    2. WebSocket event specifications
    3. Database schema with all tables and relationships
    4. Performance requirements (latency, throughput, scalability)
    5. Security requirements
    6. Integration points with other systems
    7. Error handling and edge cases
    
    Be specific, technical, and comprehensive. Use JSON format for schemas.
    """
    
    print("🤖 Generating technical specifications with Gemini...")
    tech_specs = await gemini_client.generate_text(tech_specs_prompt, temperature=0.3)
    
    # Generate detailed user journeys
    user_journey_prompt = f"""
    Generate detailed user journey flows for an In-App Customer Messaging Inbox.
    
    Include:
    1. First-time support contact journey (happy path)
    2. Returning to ongoing issue journey
    3. Multi-issue management journey
    4. Escalation from chatbot to agent journey
    5. Error scenarios and recovery flows
    
    For each journey, include:
    - User actions
    - System responses
    - UI states
    - Data flows
    - Integration points
    
    Format as detailed step-by-step flows with decision points.
    """
    
    print("🤖 Generating user journeys with Gemini...")
    user_journeys = await gemini_client.generate_text(user_journey_prompt, temperature=0.4)
    
    # Generate API specifications
    api_prompt = f"""
    Generate comprehensive REST API and WebSocket specifications for an In-App Customer Messaging Inbox.
    
    Include:
    1. All REST endpoints with HTTP methods, paths, request/response schemas
    2. Authentication and authorization
    3. WebSocket events and message formats
    4. Error codes and handling
    5. Rate limiting
    6. Pagination
    7. Filtering and sorting
    
    Be specific with JSON schemas, status codes, and examples.
    """
    
    print("🤖 Generating API specifications with Gemini...")
    api_specs = await gemini_client.generate_text(api_prompt, temperature=0.3)
    
    return {
        'tech_specs': tech_specs,
        'user_journeys': user_journeys,
        'api_specs': api_specs
    }


def add_mermaid_to_doc(doc, title, mermaid_code):
    """Add mermaid diagram section to document."""
    doc.add_heading(title, 2)
    
    # Add code block
    para = doc.add_paragraph()
    run = para.add_run(mermaid_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    note = doc.add_paragraph(style='Intense Quote')
    note.add_run("💡 Tip: Render this mermaid diagram in Notion, GitHub, or at mermaid.live").italic = True
    
    doc.add_paragraph('')


async def create_final_comprehensive_prd():
    """Create the final comprehensive PRD."""
    context = load_context_from_env()
    
    if not context.gemini_api_key:
        print("❌ Gemini API key not configured")
        return None
    
    gemini_client = GeminiClient(context.gemini_api_key, context.gemini_model_name)
    
    # Generate content with Gemini
    gemini_content = await generate_comprehensive_sections_with_gemini(gemini_client)
    
    # Create document
    doc = Document()
    
    # Title page
    title = doc.add_heading('Product Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('In-App Customer Messaging Inbox', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Version: 1.0 | Phase 1 - Q1 2026')
    doc.add_paragraph('Comprehensive Technical Specification')
    doc.add_page_break()
    
    # Add all sections with Gemini-generated content
    doc.add_heading('Technical Specifications (AI-Generated)', 1)
    doc.add_paragraph(gemini_content['tech_specs'])
    
    doc.add_page_break()
    doc.add_heading('User Journeys (AI-Generated)', 1)
    doc.add_paragraph(gemini_content['user_journeys'])
    
    doc.add_page_break()
    doc.add_heading('API Specifications (AI-Generated)', 1)
    doc.add_paragraph(gemini_content['api_specs'])
    
    # Add all mermaid diagrams
    doc.add_page_break()
    doc.add_heading('System Diagrams', 1)
    
    diagrams = {
        'System Architecture': """```mermaid
graph TB
    subgraph "Client Layer"
        A[Mobile App iOS] --> B[In-App Messaging UI]
        C[Mobile App Android] --> B
        D[Web App] --> B
    end
    
    subgraph "API Gateway"
        B --> E[Message API]
        B --> F[Thread API]
        B --> G[Notification API]
    end
    
    subgraph "Core Services"
        E --> H[Message Service]
        F --> I[Thread Management Service]
        G --> J[Push Notification Service]
        H --> K[Message Queue]
        I --> L[Thread Database]
        H --> L
    end
    
    subgraph "AI & Chatbot"
        K --> M[Chatbot Service]
        M --> N[Intent Detection]
        M --> O[Thread Titling AI]
        M --> P[Response Generation]
    end
    
    subgraph "Support Systems"
        H --> Q[Ticketing System]
        Q --> R[Ticket Status Updates]
        R --> H
        I --> S[Agent Workspace]
        S --> H
    end
    
    subgraph "Data Layer"
        L --> T[(PostgreSQL)]
        K --> U[(Redis Cache)]
        V[(Message History DB)] --> H
    end
    
    subgraph "External Services"
        J --> W[FCM/APNS]
        O --> X[Gemini AI API]
        N --> X
    end
```""",
        
        'Message Flow Sequence': """```mermaid
sequenceDiagram
    participant C as Customer
    participant UI as In-App UI
    participant API as Message API
    participant MSG as Message Service
    participant CHAT as Chatbot
    participant TICKET as Ticketing System
    participant AGENT as Support Agent
    participant NOTIF as Notification Service
    
    C->>UI: Opens Messages tab
    UI->>API: Get all threads
    API->>MSG: Fetch user threads
    MSG-->>API: Return thread list
    API-->>UI: Display threads
    UI-->>C: Show inbox
    
    C->>UI: Starts new conversation
    UI->>API: Create new thread
    API->>MSG: Initialize thread
    MSG->>CHAT: Route message
    CHAT->>CHAT: Detect intent
    CHAT->>CHAT: Generate thread title
    CHAT-->>MSG: Response + title
    MSG->>API: Thread created
    API->>NOTIF: Send notification
    API-->>UI: Display thread
    UI-->>C: Show conversation
    
    alt Chatbot resolves
        CHAT-->>C: Provide solution
    else Escalation needed
        CHAT->>TICKET: Create ticket
        TICKET->>MSG: Ticket created event
        MSG->>API: Status update
        API->>UI: Show ticket status
        UI-->>C: "Ticket created"
        TICKET->>AGENT: Assign ticket
        AGENT->>MSG: Agent response
        MSG->>API: New message
        API->>NOTIF: Push notification
        API->>UI: Update thread
        UI-->>C: Show agent message
    end
```""",
        
        'Database Schema': """```mermaid
erDiagram
    THREADS ||--o{ MESSAGES : contains
    THREADS ||--o| TICKETS : linked_to
    THREADS }o--|| USERS : belongs_to
    MESSAGES }o--|| USERS : sent_by
    MESSAGES }o--o| AGENTS : handled_by
    TICKETS ||--o{ TICKET_EVENTS : has
    THREADS }o--o{ THREAD_METADATA : has
    
    THREADS {
        uuid id PK
        uuid user_id FK
        string title
        string status
        datetime created_at
        datetime updated_at
        json metadata
        string ai_title
    }
    
    MESSAGES {
        uuid id PK
        uuid thread_id FK
        uuid user_id FK
        uuid agent_id FK
        string content
        string type
        datetime sent_at
        boolean read
        json attachments
    }
    
    TICKETS {
        uuid id PK
        uuid thread_id FK
        string status
        string priority
        datetime created_at
        datetime resolved_at
        uuid assigned_agent_id FK
    }
    
    TICKET_EVENTS {
        uuid id PK
        uuid ticket_id FK
        string event_type
        string description
        datetime occurred_at
    }
    
    USERS {
        uuid id PK
        string email
        string name
    }
    
    AGENTS {
        uuid id PK
        string name
        string email
    }
    
    THREAD_METADATA {
        uuid id PK
        uuid thread_id FK
        string intent
        json entities
        string category
    }
```""",
        
        'Thread State Machine': """```mermaid
stateDiagram-v2
    [*] --> New: Customer starts conversation
    New --> ChatbotActive: First message sent
    ChatbotActive --> ChatbotActive: Chatbot responds
    ChatbotActive --> Resolved: Issue resolved by chatbot
    ChatbotActive --> Escalated: Escalation needed
    Escalated --> TicketCreated: Ticket system creates ticket
    TicketCreated --> AgentAssigned: Agent assigned
    AgentAssigned --> InProgress: Agent responds
    InProgress --> WaitingForCustomer: Agent asks for info
    WaitingForCustomer --> InProgress: Customer responds
    InProgress --> Resolved: Issue resolved
    Resolved --> [*]
    ChatbotActive --> Archived: Customer abandons
    InProgress --> Archived: No activity for 30 days
    Archived --> [*]
```""",
        
        'Customer Journey Flow': """```mermaid
journey
    title Customer Support Journey - In-App Messaging
    section Discovery
      Customer has issue: 3: Customer
      Opens Tamara app: 4: Customer
      Navigates to Messages tab: 5: Customer
    section Initiation
      Starts new conversation: 5: Customer
      Chatbot greets and understands intent: 4: System
      AI creates thread with title: 5: System
    section Resolution
      Chatbot attempts resolution: 4: System
      Customer provides info: 3: Customer
      Issue resolved OR escalated: 5: System
    section Escalation (if needed)
      Ticket created automatically: 4: System
      Agent assigned: 4: System
      Customer sees status update: 5: Customer
      Agent responds in thread: 5: Agent
      Issue resolved: 5: Customer
    section Follow-up
      Customer returns to thread: 4: Customer
      Sees full history: 5: Customer
      Continues conversation: 4: Customer
```"""
    }
    
    for title, code in diagrams.items():
        add_mermaid_to_doc(doc, title, code)
    
    return doc


if __name__ == "__main__":
    print("=" * 80)
    print("Generating ULTIMATE Comprehensive PRD with Gemini AI")
    print("=" * 80)
    print()
    
    doc = asyncio.run(create_final_comprehensive_prd())
    
    if doc:
        output_file = Path("PRD_In-App_Messaging_ULTIMATE.docx")
        doc.save(str(output_file))
        
        print(f"\n✅ ULTIMATE Comprehensive PRD created!")
        print(f"📄 Saved to: {output_file.absolute()}")
        print(f"\nIncludes:")
        print("  ✅ AI-generated technical specifications")
        print("  ✅ AI-generated user journeys")
        print("  ✅ AI-generated API specifications")
        print("  ✅ All mermaid diagrams (architecture, sequences, state machines, ERD)")
        print("  ✅ Database schema with relationships")
        print("  ✅ Complete system documentation")
    else:
        print("❌ Failed to create PRD")



